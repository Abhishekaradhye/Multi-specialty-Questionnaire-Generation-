# models = {
#     'perplexity': {
#         'mixtral-8x7b-instruct': 'perplexity : mixtral-8x7b-instruct',
#         'llama-3.1-70b-instruct': 'perplexity : llama-3.1-70b-instruct',
#         'llama-3.1-sonar-small-128k-chat': 'perplexity : llama-3.1-sonar-small-128k-chat'
#     },
#     'openai': {
#         'gpt-3.5-turbo': 'openai : gpt-3.5-turbo',
#         'gpt-4': 'openai : gpt-4',
#         'gpt-4o-mini': 'openai : gpt-4o-mini',
#         'gpt-4-turbo': 'openai : gpt-4-turbo'
#     }
# }
import os
import openai
import requests
import csv
import pandas as pd
import time
import gradio as gr
from docx import Document

# API endpoints and keys
endpoints = {
    'perplexity': 'https://api.perplexity.ai/chat/completions',
    'openai': 'https://api.openai.com/v1/chat/completions'
}

api_keys = {
    'perplexity': 'MY_PERPLEXITY_KEY',
    'openai': 'MY_OPENAI_KEY'
}

models = {
    'perplexity': {
        'mixtral-8x7b-instruct': 'mixtral-8x7b-instruct',
        'llama-3.1-70b-instruct': 'llama-3.1-70b-instruct',
        'llama-3.1-sonar-small-128k-chat': 'llama-3.1-sonar-small-128k-chat'
    },
    'openai': {
        'gpt-3.5-turbo': 'gpt-3.5-turbo',
        'gpt-4': 'gpt-4',
        'gpt-4o-mini': 'gpt-4o-mini',
        'gpt-4-turbo': 'gpt-4-turbo'
    }
}

def call_llm_api(model_type, model, prompt, retries=3):
    for attempt in range(retries):
        try:
            if model_type == 'openai':
                response = openai.ChatCompletion.create(
                    model=model,
                    messages=[{"role": "user", "content": prompt}],
                    api_key=api_keys[model_type]
                )
                return response['choices'][0]['message']['content'].strip()
            elif model_type == 'perplexity':
                headers = {
                    'Authorization': f"Bearer {api_keys[model_type]}",
                    'Content-Type': 'application/json'
                }
                payload = {
                    'model': model,
                    'messages': [{'role': 'user', 'content': prompt}]
                }
                response = requests.post(endpoints[model_type], headers=headers, json=payload)
                return response.json()['choices'][0]['message']['content'].strip()
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 429:
                print(f"Rate limit exceeded. Retrying in {2**attempt} seconds...")
                time.sleep(2**attempt)
            else:
                print(f"HTTP error occurred: {e}")
                raise
        except Exception as e:
            print(f"An error occurred: {e}")
            raise
    raise Exception("Failed after several retries")

def process_row(row, prompt1, prompt2, folder1, folder2, model_type1, model1, model_type2, model2):
    try:
        prompt1_response = call_llm_api(model_type1, model1, prompt1)
        
        # Save Prompt 1 in Folder 1 as a text document
        with open(f"{folder1}/prompt1.txt", 'w', encoding='utf-8') as f:
            f.write(prompt1)

        reasons = prompt1_response.split(',')
        csv_filename = f"{folder1}/{row[0]} reasons.csv"
        with open(csv_filename, mode='w', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            for reason in reasons:
                writer.writerow([reason.strip()])

        for reason in reasons:
            reason = reason.strip()
            prompt2_with_reason = f"{prompt2} the reason: {reason}. For the medical specialty: {row[0]}"
            prompt2_response = call_llm_api(model_type2, model2, prompt2_with_reason)
            
            # Save Prompt 2 in Folder 2 as a text document
            with open(f"{folder2}/prompt2.txt", 'w', encoding='utf-8') as f:
                f.write(prompt2_with_reason)

            doc_filename = f"{folder2}/{row[0]}_{reason} questions.docx"
            doc = Document()
            doc.add_heading(reason, level=1)
            doc.add_paragraph(prompt2_response)
            doc.save(doc_filename)

        return f"Successfully processed row: {row[0]}"
    except Exception as e:
        return f"Error during processing row {row[0]}: {e}"

def main(userfile, folder1, folder2, prompt1_template, prompt2_template, model_type1, model1, model_type2, model2):
    # Load the userfile
    userfile_data = pd.read_csv(userfile.name)

    results = []

    for index, row in userfile_data.iterrows():
        prompt1 = f"{prompt1_template} the medical specialty: {row.iloc[0]}"
        prompt2 = f"{prompt2_template}"

        result = process_row(row, prompt1, prompt2, folder1, folder2, model_type1, model1, model_type2, model2)
        results.append(result)

    return "\n".join(results)

def run_gradio():
    def gradio_interface(userfile, folder1, folder2, prompt1, prompt2, model_type1, model1, model_type2, model2):
        os.makedirs(folder1, exist_ok=True)
        os.makedirs(folder2, exist_ok=True)
        
        return main(userfile, folder1, folder2, prompt1, prompt2, model_type1, model1, model_type2, model2)

    model_options = list(models.keys())
    openai_models = models['openai']
    perplexity_models = models['perplexity']

    with gr.Blocks() as demo:
        gr.Markdown("### CSV Processing Workflow")
        userfile = gr.File(label="Upload CSV File")
        folder1 = gr.Textbox(label="Output Folder 1", placeholder="Path to save CSV files")
        folder2 = gr.Textbox(label="Output Folder 2", placeholder="Path to save DOCX files")
        prompt1 = gr.Textbox(label="Prompt 1", placeholder="Enter the first prompt")
        prompt2 = gr.Textbox(label="Prompt 2", placeholder="Enter the second prompt")
        
        model_type1 = gr.Radio(model_options, label="Select Model Type for Process 1")
        model1 = gr.Dropdown(list(openai_models.values()), label="Select Model for Process 1")

        model_type2 = gr.Radio(model_options, label="Select Model Type for Process 2")
        model2 = gr.Dropdown(list(openai_models.values()), label="Select Model for Process 2")

        def update_model_list1(selected_model_type):
            if selected_model_type == "openai":
                return gr.update(choices=list(openai_models.values()))
            elif selected_model_type == "perplexity":
                return gr.update(choices=list(perplexity_models.values()))
        
        def update_model_list2(selected_model_type):
            if selected_model_type == "openai":
                return gr.update(choices=list(openai_models.values()))
            elif selected_model_type == "perplexity":
                return gr.update(choices=list(perplexity_models.values()))
        
        model_type1.change(fn=update_model_list1, inputs=model_type1, outputs=model1)
        model_type2.change(fn=update_model_list2, inputs=model_type2, outputs=model2)

        process_btn = gr.Button("Start Processing")
        output = gr.Textbox(label="Output")

        process_btn.click(fn=gradio_interface, inputs=[userfile, folder1, folder2, prompt1, prompt2, model_type1, model1, model_type2, model2], outputs=output)

    demo.launch()

if __name__ == "__main__":
    run_gradio()
